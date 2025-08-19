import os
import pdfplumber
from docx import Document
import subprocess
import pytesseract
from pdf2image import convert_from_path
from typing import Dict, Any, List
import logging

logger = logging.getLogger(__name__)

class FileProcessor:
    @staticmethod
    async def extract_text_from_file(file_path: str) -> str:
        """Public method to extract text from files. This is the main interface for text extraction."""
        return await FileProcessor._extract_text_from_file(file_path)
    
    @staticmethod
    async def _extract_text_from_file(file_path: str) -> str:
        try:
            if not os.path.exists(file_path):
                raise ValueError(f"File does not exist: {file_path}")

            max_file_size = 5 * 1024 * 1024
            if os.path.getsize(file_path) > max_file_size:
                raise ValueError("File size exceeds maximum limit of 5MB")

            # Validate file format before processing
            if not FileProcessor._validate_file_format(file_path):
                logger.warning(f"File format validation failed for: {file_path}")
                # Continue processing but log the warning

            file_ext = os.path.splitext(file_path)[1].lower()

            if file_ext == '.pdf':
                return await FileProcessor._extract_from_pdf(file_path)
            elif file_ext == '.docx':
                return FileProcessor._extract_from_docx(file_path)
            elif file_ext == '.doc':
                return FileProcessor._extract_from_doc(file_path)
            elif file_ext == '.txt':
                return FileProcessor._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")

        except Exception as e:
            error_msg = str(e)
            raise Exception(f"Error extracting text from {file_path}: {error_msg}")

    @staticmethod
    async def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF files with OCR fallback."""
        try:
            with open(file_path, 'rb') as f:
                # Check PDF header
                header = f.read(5)
                if header != b'%PDF-':
                    raise ValueError("Invalid PDF file: Missing PDF header")
            
            with pdfplumber.open(file_path) as pdf:
                if not pdf.pages:
                    raise ValueError("PDF file appears to be empty or corrupted")
                
                text = "\n".join(
                    page.extract_text() for page in pdf.pages
                    if page.extract_text()
                )
                
                if text.strip():
                    return text

            # If pdfplumber extracts no text, fall back to OCR
            logger.info("pdfplumber extracted no text, attempting OCR.")
            try:
                images = convert_from_path(file_path)
                ocr_text = "".join(pytesseract.image_to_string(image) for image in images)
                if not ocr_text.strip():
                    raise ValueError("OCR also failed to extract text.")
                return ocr_text
            except Exception as ocr_error:
                raise ValueError(f"Could not extract text from PDF. pdfplumber and OCR failed. OCR error: {ocr_error}")

        except Exception as e:
            logger.error(f"Initial PDF processing failed: {e}. Attempting OCR.")
            try:
                images = convert_from_path(file_path)
                ocr_text = "".join(pytesseract.image_to_string(image) for image in images)
                if not ocr_text.strip():
                    raise ValueError("OCR also failed to extract text.")
                return ocr_text
            except Exception as ocr_error:
                raise ValueError(f"Error processing PDF file: {e}, and OCR fallback also failed: {ocr_error}")

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX files."""
        try:
            # First validate it's actually a DOCX file
            with open(file_path, 'rb') as f:
                header = f.read(4)
                if header != b'PK\x03\x04' and header != b'PK\x05\x06' and header != b'PK\x07\x08':
                    raise ValueError("File does not appear to be a valid DOCX (ZIP) file")

            doc = Document(file_path)
            text_parts = [p.text for p in doc.paragraphs]
            
            # Extract text from tables as well
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            
            text = "\n".join(text_parts)
            if not text.strip():
                raise ValueError("Could not extract text from document. The file might be corrupted or empty")
            return text
        except Exception as e:
            # Check if it's actually a .doc file with wrong extension
            if "not a Word file" in str(e) or "themeManager" in str(e):
                logger.warning(f"File appears to be .doc format with .docx extension: {file_path}")
                return FileProcessor._extract_from_doc(file_path)
            raise ValueError(f"Error processing .docx file: {str(e)}")

    @staticmethod
    def _extract_from_doc(file_path: str) -> str:
        """Extract text from legacy DOC files using antiword.

        Gracefully handle cases where a .doc file is actually plain text (e.g. a test file)
        or where the file was misnamed. Attempt to read as text with common encodings
        before invoking antiword. If the file is actually a DOCX with a .doc extension,
        delegate to the DOCX extractor.
        """
        try:
            # First check if it's actually a DOC file
            with open(file_path, 'rb') as f:
                header = f.read(8)
                # If it's a DOCX with .doc extension, delegate
                if header.startswith(b'PK'):
                    logger.warning("File appears to be DOCX format with .doc extension")
                    return FileProcessor._extract_from_docx(file_path)

                # If header doesn't match OLE2 signatures, attempt a plain-text fallback
                ole_signatures = (b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1', b'\xdb\xa5-\x00\x00\x00')
                if not any(header.startswith(sig) for sig in ole_signatures):
                    logger.warning("DOC file header not recognized as OLE2. Trying plain-text fallback before antiword.")

                    # Try reading as text with common encodings
                    encodings = ['utf-8', 'cp1252', 'latin-1', 'utf-16']
                    for enc in encodings:
                        try:
                            with open(file_path, 'r', encoding=enc) as tf:
                                text = tf.read()
                                if text and text.strip():
                                    logger.info(f"Successfully read .doc file as plain text using encoding: {enc}")
                                    return text
                        except Exception:
                            # Try next encoding
                            continue
                    # If plain-text fallback didn't return, continue to try antiword below

            # Use antiword for real legacy .doc files
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                check=True,
                timeout=30
            )
            text = result.stdout
            if not text.strip():
                raise ValueError("antiword failed to extract text or file is empty")
            return text

        except FileNotFoundError:
            raise ValueError(
                "antiword command not found. Please ensure it is installed and in the system's PATH. "
                "Alternative: Convert the .doc file to .docx or .pdf format."
            )
        except subprocess.CalledProcessError as e:
            error_detail = e.stderr if e.stderr else "Unknown antiword error"
            raise ValueError(f"Error processing .doc file with antiword: {error_detail}")
        except subprocess.TimeoutExpired:
            raise ValueError("Processing .doc file timed out. File might be too large or corrupted.")
        except Exception as e:
            raise ValueError(f"An unexpected error occurred while processing the .doc file: {str(e)}")

    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Extract text from TXT files with encoding detection."""
        try:
            # Try UTF-8 first (most common)
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                if not text.strip():
                    raise ValueError("Text file is empty")
                return text
        except UnicodeDecodeError:
            # Try other common encodings
            encodings = ['latin-1', 'cp1252', 'utf-16', 'ascii']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                        if text.strip():
                            logger.info(f"Successfully read text file with {encoding} encoding")
                            return text
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            # Last resort with error replacement
            try:
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
                    if text.strip():
                        logger.warning("Text file read with character replacement due to encoding issues")
                        return text
                    else:
                        raise ValueError("Text file is empty")
            except Exception as e:
                raise ValueError(f"Could not read text file with any encoding: {str(e)}")
        except Exception as e:
            raise ValueError(f"Error reading text file: {str(e)}")

    @staticmethod
    def _validate_file_format(file_path: str) -> bool:
        """Validate if the file format matches its extension."""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)
            
            if file_ext == '.pdf':
                return header.startswith(b'%PDF')
            elif file_ext == '.docx':
                # DOCX files are ZIP archives starting with PK
                return header.startswith(b'PK')
            elif file_ext == '.doc':
                # DOC files have OLE2 signature or Word 6.0/95 signature
                return (header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1') or  # OLE2
                        header.startswith(b'\xdb\xa5-\x00\x00\x00'))  # Word 6.0/95
            elif file_ext == '.txt':
                # For TXT, just check if it's readable
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        f.read(100)  # Try to read first 100 chars
                    return True
                except UnicodeDecodeError:
                    return True  # Still valid, just different encoding
            
        except Exception as e:
            logger.error(f"File validation failed: {str(e)}")
            return False
        
        return True

    @staticmethod
    def get_file_info(file_path: str) -> Dict[str, Any]:
        """Get detailed file information for debugging."""
        info = {
            'exists': os.path.exists(file_path),
            'size': 0,
            'extension': '',
            'content_type_guess': '',
            'header_bytes': b'',
            'is_valid_format': False
        }
        
        if os.path.exists(file_path):
            info['size'] = os.path.getsize(file_path)
            info['extension'] = os.path.splitext(file_path)[1].lower()
            
            try:
                with open(file_path, 'rb') as f:
                    info['header_bytes'] = f.read(16)
                
                # Guess content type from header
                header = info['header_bytes']
                if header.startswith(b'%PDF'):
                    info['content_type_guess'] = 'application/pdf'
                elif header.startswith(b'PK'):
                    info['content_type_guess'] = 'application/zip (likely DOCX)'
                elif header.startswith(b'\xd0\xcf\x11\xe0'):
                    info['content_type_guess'] = 'application/msword (DOC)'
                else:
                    info['content_type_guess'] = 'unknown'
                
                info['is_valid_format'] = FileProcessor._validate_file_format(file_path)
                
            except Exception as e:
                logger.error(f"Error getting file info: {str(e)}")
        
        return info